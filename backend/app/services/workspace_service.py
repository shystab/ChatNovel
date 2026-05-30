from __future__ import annotations

import io
import json
import re
import zipfile
from datetime import datetime
from html import escape
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable
from urllib.parse import quote, unquote

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from sqlmodel import Session, select

from app.core.config import settings
from app.models.books import Book
from app.models.chapters import Chapter
from app.models.setting import Setting


INVALID_FILENAME_CHARS = r'<>:"/\|?*'
CHAPTER_FILENAME_RE = re.compile(r"^(?P<order>\d+)-(?P<id>\d+)-(?P<title>.+)\.txt$")


class _PlainTextHTMLParser(HTMLParser):
    block_tags = {"p", "div", "br", "li", "h1", "h2", "h3", "blockquote"}

    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag == "br":
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self.block_tags:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        self.parts.append(data)

    def text(self) -> str:
        raw = "".join(self.parts)
        lines = [line.strip() for line in raw.splitlines()]
        output: list[str] = []
        previous_blank = False
        for line in lines:
            if not line:
                if not previous_blank:
                    output.append("")
                previous_blank = True
                continue
            output.append(line)
            previous_blank = False
        return "\n".join(output).strip()


def html_to_plain_text(content: str) -> str:
    if not content:
        return ""
    if "<" not in content or ">" not in content:
        return content.strip()
    parser = _PlainTextHTMLParser()
    parser.feed(content)
    return parser.text()


def normalize_novel_text(content: str) -> str:
    text = html_to_plain_text(content)
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n|\n", text) if p.strip()]
    return "\n\n".join(paragraphs)


def plain_text_to_html(content: str) -> str:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n|\n", content) if p.strip()]
    if not paragraphs:
        return ""
    return "".join(f"<p>{escape(paragraph)}</p>" for paragraph in paragraphs)


def safe_filename(value: str, fallback: str = "untitled") -> str:
    cleaned = value.strip() or fallback
    for char in INVALID_FILENAME_CHARS:
        cleaned = cleaned.replace(char, "_")
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    return cleaned[:80] or fallback


def content_disposition(filename: str) -> str:
    ascii_filename = filename.encode("ascii", "ignore").decode().strip() or "download"
    if ascii_filename.startswith("."):
        ascii_filename = f"download{ascii_filename}"
    ascii_filename = ascii_filename.replace("\\", "_").replace('"', "_")
    quoted_filename = quote(filename)
    return f"attachment; filename=\"{ascii_filename}\"; filename*=UTF-8''{quoted_filename}"


def workspace_root() -> Path:
    root = Path(_configured_workspace_dir()).expanduser()
    if not root.is_absolute():
        root = Path.cwd() / root
    root.mkdir(parents=True, exist_ok=True)
    return root


def _configured_workspace_dir() -> str:
    try:
        from app.db.session import engine

        with Session(engine) as session:
            db_settings = session.exec(select(Setting)).first()
            if db_settings and db_settings.workspace_dir:
                return db_settings.workspace_dir
    except Exception:
        pass
    return settings.NOVEL_WORKSPACE_DIR


def book_folder(book: Book | None, book_id: int | None = None) -> Path:
    if book:
        name = f"{book.id:03d}-{safe_filename(book.title, 'book')}"
    elif book_id:
        name = f"{book_id:03d}-book"
    else:
        name = "000-unassigned"
    folder = workspace_root() / name
    (folder / "chapters").mkdir(parents=True, exist_ok=True)
    (folder / "exports").mkdir(parents=True, exist_ok=True)
    (folder / ".cache").mkdir(parents=True, exist_ok=True)
    return folder


def write_project_manifest(book: Book) -> Path:
    folder = book_folder(book)
    manifest = {
        "id": book.id,
        "title": book.title,
        "description": book.description,
        "user_id": book.user_id,
        "updated_at": datetime.now().isoformat(timespec="seconds"),
    }
    path = folder / "project.json"
    path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def write_library_index(items: Iterable[tuple[Book, list[Chapter]]]) -> dict:
    item_list = list(items)
    root = workspace_root()
    updated_at = datetime.now().isoformat(timespec="seconds")
    books: list[dict] = []
    readme_lines = [
        "# VibeWriter 作品库",
        "",
        f"更新时间：{updated_at}",
        "",
        "## 作品",
        "",
    ]

    for book, chapters in item_list:
        folder = book_folder(book)
        relative_folder = folder.relative_to(root).as_posix()
        char_count = sum(len(normalize_novel_text(chapter.content)) for chapter in chapters)
        book_info = {
            "id": book.id,
            "title": book.title,
            "description": book.description,
            "folder": relative_folder,
            "chapter_count": len(chapters),
            "char_count": char_count,
            "updated_at": book.update_time.isoformat() if book.update_time else None,
        }
        books.append(book_info)
        readme_lines.append(
            f"- [{book.title}]({relative_folder}/project.json)：{len(chapters)} 章，{char_count} 字符"
        )

    if not books:
        readme_lines.append("- 暂无作品")

    manifest = {
        "app": "VibeWriter",
        "updated_at": updated_at,
        "book_count": len(books),
        "chapter_count": sum(book["chapter_count"] for book in books),
        "books": books,
    }
    manifest_path = root / "library.json"
    readme_path = root / "README.md"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    readme_path.write_text("\n".join(readme_lines).rstrip() + "\n", encoding="utf-8")
    return {
        "workspace": str(root),
        "manifest": str(manifest_path),
        "readme": str(readme_path),
        **manifest,
    }


def scan_library_workspace() -> dict:
    root = workspace_root()
    books: list[dict] = []

    for folder in sorted(path for path in root.iterdir() if path.is_dir()):
        if folder.name.startswith("."):
            continue

        manifest_path = folder / "project.json"
        manifest: dict = {}
        if manifest_path.exists():
            try:
                manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            except json.JSONDecodeError:
                manifest = {}

        title = manifest.get("title") or re.sub(r"^\d+-", "", folder.name)
        chapters: list[dict] = []
        chapters_dir = folder / "chapters"
        if chapters_dir.exists():
            for index, path in enumerate(sorted(chapters_dir.glob("*.txt")), start=1):
                match = CHAPTER_FILENAME_RE.match(path.name)
                order = int(match.group("order")) if match else index
                chapter_id = int(match.group("id")) if match else None
                chapter_title = match.group("title") if match else path.stem
                text = path.read_text(encoding="utf-8")
                chapters.append({
                    "id": chapter_id,
                    "title": chapter_title,
                    "order": order,
                    "path": str(path),
                    "char_count": len(text.strip()),
                })

        books.append({
            "id": manifest.get("id"),
            "title": title,
            "description": manifest.get("description"),
            "user_id": manifest.get("user_id") or "default_user",
            "folder": str(folder),
            "chapter_count": len(chapters),
            "char_count": sum(chapter["char_count"] for chapter in chapters),
            "chapters": chapters,
        })

    return {
        "workspace": str(root),
        "book_count": len(books),
        "chapter_count": sum(book["chapter_count"] for book in books),
        "char_count": sum(book["char_count"] for book in books),
        "books": books,
    }


def read_workspace_book(folder: str) -> dict:
    root = workspace_root().resolve()
    book_folder_path = Path(folder).resolve()
    if root not in book_folder_path.parents and book_folder_path != root:
        raise ValueError("Book folder must be inside workspace")

    manifest_path = book_folder_path / "project.json"
    manifest: dict = {}
    if manifest_path.exists():
        try:
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            manifest = {}

    title = manifest.get("title") or re.sub(r"^\d+-", "", book_folder_path.name)
    chapters: list[dict] = []
    chapters_dir = book_folder_path / "chapters"
    if chapters_dir.exists():
        for index, path in enumerate(sorted(chapters_dir.glob("*.txt")), start=1):
            match = CHAPTER_FILENAME_RE.match(path.name)
            order = int(match.group("order")) if match else index
            chapter_id = int(match.group("id")) if match else None
            chapter_title = match.group("title") if match else path.stem
            text = path.read_text(encoding="utf-8")
            chapters.append({
                "id": chapter_id,
                "title": chapter_title,
                "order": order,
                "content": plain_text_to_html(text),
            })

    return {
        "id": manifest.get("id"),
        "title": title,
        "description": manifest.get("description"),
        "user_id": manifest.get("user_id") or "default_user",
        "chapters": chapters,
    }


def chapter_file_path(chapter: Chapter, book: Book | None = None) -> Path:
    folder = book_folder(book, chapter.book_id)
    order = chapter.order or chapter.id or 0
    chapter_id = chapter.id or 0
    filename = f"{order:03d}-{chapter_id:04d}-{safe_filename(chapter.title, 'chapter')}.txt"
    return folder / "chapters" / filename


def chapter_files(chapter: Chapter) -> list[Path]:
    if not chapter.id:
        return []
    pattern = f"*-{chapter.id:04d}-*.txt"
    return [path for path in workspace_root().glob(f"*/chapters/{pattern}") if path.is_file()]


def delete_chapter_files(chapter: Chapter) -> list[Path]:
    deleted: list[Path] = []
    for path in chapter_files(chapter):
        path.unlink(missing_ok=True)
        deleted.append(path)
    return deleted


def write_chapter_file(chapter: Chapter, book: Book | None = None) -> Path:
    path = chapter_file_path(chapter, book)
    target = path.resolve()
    for stale_path in chapter_files(chapter):
        if stale_path.resolve() != target:
            stale_path.unlink(missing_ok=True)
    text = normalize_novel_text(chapter.content)
    path.write_text(text + ("\n" if text else ""), encoding="utf-8")
    return path


def build_txt_export(chapters: Iterable[Chapter], title: str = "小说全集") -> str:
    sections = [title.strip() or "小说全集", ""]
    for chapter in chapters:
        body = normalize_novel_text(chapter.content)
        sections.append(chapter.title.strip() or f"第 {chapter.order} 章")
        sections.append("")
        if body:
            sections.append(body)
            sections.append("")
    return "\n".join(sections).rstrip() + "\n"


def build_docx_export(chapters: Iterable[Chapter], title: str = "小说全集") -> io.BytesIO:
    chapter_list = list(chapters)
    doc = docx.Document()
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.75

    title_paragraph = doc.add_heading(title.strip() or "小说全集", 0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for index, chapter in enumerate(chapter_list):
        heading = doc.add_heading(chapter.title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        body = normalize_novel_text(chapter.content)
        for paragraph_text in body.split("\n\n"):
            if not paragraph_text.strip():
                continue
            paragraph = doc.add_paragraph(paragraph_text.strip())
            paragraph.paragraph_format.first_line_indent = Pt(24)
            paragraph.paragraph_format.line_spacing = 1.75
        if index != len(chapter_list) - 1:
            doc.add_page_break()

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def sync_book_workspace(book: Book, chapters: Iterable[Chapter]) -> dict:
    folder = book_folder(book)
    manifest_path = write_project_manifest(book)
    written = [write_chapter_file(chapter, book) for chapter in chapters]
    return {
        "workspace": str(folder),
        "manifest": str(manifest_path),
        "chapter_count": len(written),
        "chapters": [str(path) for path in written],
    }


def sync_library_workspace(items: Iterable[tuple[Book, list[Chapter]]]) -> dict:
    item_list = list(items)
    synced_books = [sync_book_workspace(book, chapters) for book, chapters in item_list]
    index = write_library_index(item_list)
    return {
        **index,
        "synced_books": synced_books,
    }


def _sqlite_database_path(database_url: str) -> Path | None:
    if not database_url.startswith("sqlite"):
        return None
    if database_url == "sqlite:///:memory:":
        return None

    if database_url.startswith("sqlite:////"):
        raw_path = "/" + database_url.removeprefix("sqlite:////")
    elif database_url.startswith("sqlite:///"):
        raw_path = database_url.removeprefix("sqlite:///")
    else:
        return None

    path = Path(unquote(raw_path))
    if not path.is_absolute():
        path = Path.cwd() / path
    return path if path.exists() else None


def build_workspace_backup(database_url: str | None = None) -> tuple[str, io.BytesIO]:
    root = workspace_root()
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    filename = f"VibeWriter-backup-{timestamp}.zip"
    database_path = _sqlite_database_path(database_url or settings.DATABASE_URL)

    metadata = {
        "app": "VibeWriter",
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "workspace": str(root),
        "database": str(database_path) if database_path else None,
    }

    output = io.BytesIO()
    with zipfile.ZipFile(output, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("backup.json", json.dumps(metadata, ensure_ascii=False, indent=2))

        if database_path:
            archive.write(database_path, f"database/{database_path.name}")

        for path in root.rglob("*"):
            if not path.is_file():
                continue
            archive.write(path, f"workspace/{path.relative_to(root).as_posix()}")

    output.seek(0)
    return filename, output
